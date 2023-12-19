import gcsfs
from PIL import Image
import streamlit as st
import joblib
from pathlib import Path
from pipeline_helpers import ExtractAndPredict, CleanParagraphs, TitleGenerator
from predictors import search_paragraphs
from utils import convert_to_html
from config import UPLOAD_FOLDER
import uuid
import docx
from docx.shared import Inches
import base64
import os
import docx
import html2text


def get_binary_file_downloader_html(bin_file, file_label='File'):
    with open(bin_file, 'rb') as f:
        data = f.read()
    bin_str = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{os.path.basename(bin_file)}">Download {file_label}</a>'
    return href


# Define the Google Drive link
gcs_path = "https://drive.google.com/file/d/1-3XumxPf1PMlTdEJ4YEEavHII-bs56dJ/view?usp=drive_link"

# Load the pickle file using Joblib and gcsfs
pipeline = joblib.load(gcsfs.GCSFileSystem().open(gcs_path, "rb"))


st.title("Corportica Projectica 💼")

uploaded_file = st.file_uploader(
    "Choose a file", type=['docx', 'txt', 'pdf'], accept_multiple_files=False)

if uploaded_file is not None:
    file_name = uploaded_file.name
    file_path = Path(f"uploaded_files/{file_name}")

    # Save file to upload folder only if it doesn't exist
    if not file_path.exists():
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

    # Cache the pipeline function
    @st.cache_resource
    def run_pipeline(file_path):
        return pipeline.transform(str(file_path))

    # Cache the html conversion function
    @st.cache_resource
    def convert_to_html_file(paragraphs, output_file_name):
        html_path = Path(f"output_files/{output_file_name}.html")
        convert_to_html(paragraphs, html_path)
        return html_path

    # Run pipeline and html conversion only once
    paragraphs = run_pipeline(file_path)
    output_file_name = file_name.split('.')[0] + "_" + uuid.uuid4().hex
    html_path = convert_to_html_file(paragraphs, output_file_name)

    # Display output file
    st.write("Processed Document: ")
    with open(html_path) as f:
        contents = f.read()
    st.components.v1.html(contents, width=700, height=800, scrolling=True)
    # Add download button

    def get_docx(html_content):
        # Convert HTML to Markdown
        h = html2text.HTML2Text()
        h.ignore_links = False
        markdown_content = h.handle(html_content)

        # Create a new Word document
        doc = docx.Document()
        doc.add_heading("Processed Output", 0)

        # Adding Markdown content to the Word document
        for line in markdown_content.split('\n'):
            if line.startswith('#'):  # Heading in Markdown
                level = line.count('#')
                doc.add_heading(line.replace('#', '').strip(), level=level)
            else:
                doc.add_paragraph(line)

        return doc
    docx_doc = get_docx(contents)

    @st.cache_resource
    def get_table_download_link(_docx_doc):
        filename = f'{file_name}.docx'
        _docx_doc.save(filename)
        return filename

    docx_doc = get_docx(contents)

    docx_filename = get_table_download_link(docx_doc)
    st.markdown(get_binary_file_downloader_html(
        docx_filename, 'Word Document'), unsafe_allow_html=True)
# Search functionality outside of the file upload block
query = st.text_input("Search query:")

if query and uploaded_file is not None:

    results = search_paragraphs(query, paragraphs)

    for result in results:
        similarity_score, idx, title, content = result

        # Display divider
        st.markdown("---")

        # Display paragraph number
        st.subheader(f"Paragraph {idx+1}")

        # Display title
        st.markdown(f"**{title}**")

        # Display similarity score
        st.text(f"Similarity Score: {similarity_score:.2f}")

        # Display paragraph content
        st.write(content)
